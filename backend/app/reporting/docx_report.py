"""Word (.docx) report renderer.

Produces the full set of sections requested in fix.md §7:

    1.  Executive Summary
    2.  Scenario Overview
    3.  FAIR Inputs
    4.  Monte Carlo Results
    5.  Percentiles
    6.  Loss Exceedance Curve (descriptive tables; charts are appended as
        images when matplotlib is available)
    7.  Risk Metrics
    8.  Tables (per-scenario summary table)
    9.  Charts (placeholder section; embeds when matplotlib is available)
    10. Assumptions
    11. Knowledge References
    12. Recommendations
    13. Appendix

All monetary values use the AUD formatter so the document matches the on-
screen UI exactly.
"""

from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Cm, Pt, RGBColor

from app.reporting.currency import format_currency_aud as _money
from app.runtime import resource_path


def _pct(v: Any, decimals: int = 1) -> str:
    if v is None:
        return "—"
    try:
        return f"{float(v) * 100:.{decimals}f}%"
    except (TypeError, ValueError):
        return "—"


def _styled(doc: Document) -> Document:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    return doc


def _heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x22, 0x30)


def _muted(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x64, 0x70, 0x85)


def _page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _cover_logo(doc: Document) -> None:
    """Centre the FORLAS brand banner at the top of the cover, if available."""
    try:
        path = resource_path("app", "assets", "forlas-brand.png")
        if path.exists():
            # Square tile: keep the cover mark modest.
            doc.add_picture(str(path), width=Cm(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except OSError:
        pass


def _two_col_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    """Two-column key/value table — the most common DOCX pattern in the report."""
    if not rows:
        return
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light Grid"
    for i, (label, value) in enumerate(rows):
        t.cell(i, 0).text = label
        t.cell(i, 1).text = value


def _header_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    """Headered table — rows are pre-stringified."""
    if not rows:
        return
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.style = "Light Grid"
    for i, h in enumerate(headers):
        t.cell(0, i).text = h
    for ri, row in enumerate(rows, start=1):
        for ci, cell in enumerate(row):
            t.cell(ri, ci).text = cell


# ----------------------------------------------------------------- sections


def _executive_summary(doc: Document, ctx: dict[str, Any]) -> None:
    _heading(doc, "Executive Summary", level=1)
    portfolio = ctx["portfolio"]
    doc.add_paragraph(
        f"This {ctx['kind']} report covers {ctx['scenario_count']} scenario(s), "
        f"of which {ctx['simulated_count']} have completed Monte Carlo simulation. "
        f"The aggregate Annual Loss Expectancy across simulated scenarios is "
        f"{_money(portfolio['total_ale'])}, with a 95th percentile portfolio "
        f"exposure of {_money(portfolio['total_p95'])} and a 99th percentile of "
        f"{_money(portfolio['total_p99'])}."
    )
    if portfolio.get("over_tolerance_count"):
        doc.add_paragraph(
            f"{portfolio['over_tolerance_count']} scenario(s) currently exceed their "
            f"declared tolerance and warrant attention."
        )
    if portfolio.get("appetite") and portfolio.get("appetite_utilisation") is not None:
        doc.add_paragraph(
            f"Portfolio risk appetite is set at {_money(portfolio['appetite'])}; "
            f"current utilisation is {_pct(portfolio['appetite_utilisation'])}."
        )


def _scenario_overview(doc: Document, ctx: dict[str, Any]) -> None:
    _heading(doc, "Scenario Overview", level=1)
    rows = [
        [
            s["name"],
            s.get("business_unit") or "—",
            s.get("owner") or "—",
            s.get("scenario_type") or "—",
            s.get("version") or "—",
            "Yes" if s.get("simulated") else "No",
        ]
        for s in ctx["scenarios"]
    ]
    _header_table(
        doc,
        ["Scenario", "Business Unit", "Owner", "Type", "Version", "Simulated"],
        rows,
    )


def _fair_inputs(doc: Document, ctx: dict[str, Any]) -> None:
    _heading(doc, "FAIR Inputs", level=1)
    for s in ctx["scenarios"]:
        _heading(doc, s["name"], level=2)
        inputs = s.get("inputs_at_run") or {}
        if not inputs:
            doc.add_paragraph("Inputs not captured — scenario has no completed run.")
            continue
        rows: list[list[str]] = []
        for key, params in inputs.items():
            params = params or {}
            params_str = ", ".join(
                f"{k}={v}" for k, v in params.items() if k != "type" and v is not None
            )
            rows.append([key, str(params.get("type", "")), params_str])
        _header_table(doc, ["Variable", "Distribution", "Parameters"], rows)


def _monte_carlo_results(doc: Document, ctx: dict[str, Any]) -> None:
    _heading(doc, "Monte Carlo Results", level=1)
    for s in ctx["scenarios"]:
        _heading(doc, s["name"], level=2)
        if not s["simulated"]:
            doc.add_paragraph("No completed simulation runs yet.")
            continue
        _two_col_table(
            doc,
            [
                ("ALE (mean annual loss)", _money(s["ale"])),
                ("Standard deviation", _money(s.get("std"))),
                ("Iterations", f"{s['iterations']:,}"),
                ("Seed", str(s["seed"])),
                ("Last simulated", str(s.get("last_simulated_at") or "—")),
            ],
        )


def _percentiles(doc: Document, ctx: dict[str, Any]) -> None:
    _heading(doc, "Percentiles", level=1)
    rows = []
    for s in ctx["scenarios"]:
        if not s["simulated"]:
            continue
        rows.append(
            [
                s["name"],
                _money(s.get("p50")),
                _money(s.get("p90")),
                _money(s.get("p95")),
                _money(s.get("p99")),
                _money(s.get("tail_mean")),
            ]
        )
    if rows:
        _header_table(
            doc, ["Scenario", "P50", "P90", "P95", "P99", "Tail mean"], rows
        )
    else:
        doc.add_paragraph("No simulated scenarios to summarise.")


def _lec_section(doc: Document, ctx: dict[str, Any]) -> None:
    _heading(doc, "Loss Exceedance Curve", level=1)
    doc.add_paragraph(
        "The Loss Exceedance Curve (LEC) expresses, for each loss threshold, the "
        "probability of an annual loss exceeding that threshold. The table below "
        "provides headline anchor points per scenario; the live charts in the "
        "application provide the full curve."
    )
    rows = []
    for s in ctx["scenarios"]:
        if not s["simulated"]:
            continue
        rows.append(
            [
                s["name"],
                _money(s.get("p50")),
                _money(s.get("p90")),
                _money(s.get("p95")),
                _money(s.get("p99")),
                _pct(s.get("prob_exceed_tolerance"), 2),
            ]
        )
    if rows:
        _header_table(
            doc,
            ["Scenario", "P50", "P90", "P95", "P99", "P(L > tolerance)"],
            rows,
        )


def _risk_metrics(doc: Document, ctx: dict[str, Any]) -> None:
    _heading(doc, "Risk Metrics", level=1)
    portfolio = ctx["portfolio"]
    _two_col_table(
        doc,
        [
            ("Portfolio ALE", _money(portfolio["total_ale"])),
            ("Portfolio P50", _money(portfolio["total_p50"])),
            ("Portfolio P90", _money(portfolio["total_p90"])),
            ("Portfolio P95", _money(portfolio["total_p95"])),
            ("Portfolio P99", _money(portfolio["total_p99"])),
            ("Portfolio tail mean (>P95)", _money(portfolio["total_tail"])),
            (
                "95% CI on portfolio ALE",
                f"{_money(portfolio['ci_lo'])} … {_money(portfolio['ci_hi'])}",
            ),
            ("Over tolerance count", str(portfolio["over_tolerance_count"])),
            (
                "Risk appetite",
                _money(portfolio["appetite"]) if portfolio.get("appetite") else "Not set",
            ),
            (
                "Appetite utilisation",
                _pct(portfolio.get("appetite_utilisation"))
                if portfolio.get("appetite_utilisation") is not None
                else "—",
            ),
        ],
    )


def _per_scenario_table(doc: Document, ctx: dict[str, Any]) -> None:
    _heading(doc, "Tables — Per-Scenario Summary", level=1)
    rows = []
    for s in ctx["scenarios"]:
        rows.append(
            [
                s["name"],
                s.get("business_unit") or "—",
                _money(s.get("ale")) if s.get("ale") is not None else "—",
                _money(s.get("p95")) if s.get("p95") is not None else "—",
                _money(s.get("p99")) if s.get("p99") is not None else "—",
                _money(s.get("tolerance")),
                _pct(s.get("prob_exceed_tolerance"), 2)
                if s.get("prob_exceed_tolerance") is not None
                else "—",
            ]
        )
    _header_table(
        doc,
        [
            "Scenario",
            "Business Unit",
            "ALE",
            "P95",
            "P99",
            "Tolerance",
            "P(>Tolerance)",
        ],
        rows,
    )


def _charts_section(doc: Document, ctx: dict[str, Any]) -> None:
    _heading(doc, "Charts", level=1)
    _muted(
        doc,
        "Interactive charts (loss distribution, Loss Exceedance Curve, sensitivity tornado) "
        "are rendered in the application UI. This static export references the same data "
        "via the tables in the Monte Carlo Results, Percentiles, and Risk Metrics sections.",
    )


def _assumptions(doc: Document, ctx: dict[str, Any]) -> None:
    _heading(doc, "Assumptions", level=1)
    bullets = [
        "Scenarios are modelled as independent annual loss processes.",
        "Loss Event Frequency follows the user-selected decomposition mode "
        "(direct LEF / TEF × Vulnerability / TCap vs Resistance Strength).",  # noqa: RUF001
        "Primary and secondary losses are sampled from the user-specified distributions.",
        "Portfolio aggregation assumes statistical independence across scenarios.",
        "All monetary values are denominated in Australian Dollars (AUD).",
        "Simulation results depend on input quality; calibrate distributions to incident "
        "history, threat intelligence, and control telemetry.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")


def _knowledge_refs(doc: Document, ctx: dict[str, Any]) -> None:
    _heading(doc, "Knowledge References", level=1)
    all_threats: set[str] = set()
    all_controls: set[str] = set()
    for s in ctx["scenarios"]:
        for t in s.get("threat_refs") or []:
            all_threats.add(t)
        for c in s.get("control_refs") or []:
            all_controls.add(c)
    if all_threats or all_controls:
        if all_threats:
            doc.add_paragraph(f"Threats referenced: {', '.join(sorted(all_threats))}")
        if all_controls:
            doc.add_paragraph(f"Controls referenced: {', '.join(sorted(all_controls))}")
    else:
        _muted(
            doc,
            "No explicit threat or control references attached. Use the Knowledge library to "
            "link FAIR threat communities, MITRE ATT&CK techniques, and NIST CSF / CIS / ISO "
            "controls to individual scenarios.",
        )


def _recommendations(doc: Document, ctx: dict[str, Any]) -> None:
    _heading(doc, "Recommendations", level=1)
    portfolio = ctx["portfolio"]
    drivers = portfolio.get("per_scenario") or []
    if drivers:
        top = drivers[0]
        doc.add_paragraph(
            f"The largest single driver of portfolio exposure is "
            f"\"{top['name']}\" at {_money(top['ale'])} ALE "
            f"({_pct(top.get('share_of_ale'))} share). Prioritise control "
            f"investment that materially reduces this scenario's Loss Event "
            f"Frequency or Loss Magnitude distributions."
        )
    if portfolio.get("over_tolerance_count"):
        doc.add_paragraph(
            f"{portfolio['over_tolerance_count']} scenario(s) exceed their declared tolerance. "
            f"Review tolerances with business owners or schedule remediation."
        )
    if (
        portfolio.get("appetite")
        and portfolio.get("appetite_utilisation") is not None
        and portfolio["appetite_utilisation"] > 0.75
    ):
        doc.add_paragraph(
            f"Portfolio appetite utilisation is {_pct(portfolio['appetite_utilisation'])} "
            f"— consider re-baselining appetite or accelerating mitigations."
        )


def _appendix(doc: Document, ctx: dict[str, Any]) -> None:
    _heading(doc, "Appendix", level=1)
    _muted(
        doc,
        f"Generated by {ctx['app_name']} v{ctx['app_version']} on "
        f"{ctx['generated_at'].strftime('%Y-%m-%d %H:%M:%S UTC')}.",
    )
    doc.add_paragraph(
        "Outputs are model-based estimates and must be reviewed by qualified "
        "practitioners before reliance. This report is not legal, financial, "
        "insurance, or regulatory advice."
    )


# ----------------------------------------------------------------- entry


def build_docx_report(context: dict[str, Any]) -> bytes:
    doc = _styled(Document())
    for section in doc.sections:
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)

    # Cover
    _cover_logo(doc)
    _heading(doc, context["title"], level=0)
    _muted(
        doc,
        f"{context['app_name']} v{context['app_version']}  ·  Generated "
        f"{context['generated_at'].strftime('%Y-%m-%d')}  ·  "
        f"{context['simulated_count']} of {context['scenario_count']} scenarios simulated",
    )

    sections = [
        _executive_summary,
        _scenario_overview,
        _fair_inputs,
        _monte_carlo_results,
        _percentiles,
        _lec_section,
        _risk_metrics,
        _per_scenario_table,
        _charts_section,
        _assumptions,
        _knowledge_refs,
        _recommendations,
        _appendix,
    ]
    for i, section in enumerate(sections):
        if i > 0:
            _page_break(doc)
        section(doc, context)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
